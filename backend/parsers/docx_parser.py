from docx import Document

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        
        print(f"[DOCX Parser] Opening DOCX with {len(doc.paragraphs)} paragraphs")
        
        paragraphs = []
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(p.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        text = "\n".join(paragraphs)
        print(f"[DOCX Parser] Total extracted: {len(text)} characters")
        
        return text.strip()
    except Exception as e:
        print(f"[DOCX Parser] Error: {e}")
        import traceback
        traceback.print_exc()
        return ""

